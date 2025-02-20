# Standard Library Imports
import os, io
from datetime import datetime, timezone

# Flask & Extensions
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
#from flask_migrate import Migrate
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_mail import Mail
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor

# Models (Database Tables)
from models import db, User, Resume

# Utility Functions
from utils.s3_helper import upload_to_s3, generate_presigned_url
from utils.ai_helper import analyze_resume, generate_interview_question, evaluate_star_response, generate_resume, fine_tuner
from email_sender import send_email
from token_utils import generate_timed_token, confirm_timed_token

# File Handling
import docx
import PyPDF2
import logging


# Load environment variables
load_dotenv()

DATABASE_URL = os.getenv("DATABASE_URL")

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY')
app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['AWS_ACCESS_KEY_ID'] = os.getenv('AWS_ACCESS_KEY_ID')
app.config['AWS_SECRET_ACCESS_KEY'] = os.getenv('AWS_SECRET_ACCESS_KEY')
app.config['S3_BUCKET_NAME'] = os.getenv('S3_BUCKET_NAME')
app.config['MAIL_DEBUG'] = True
app.config["MAIL_SERVER"] = os.getenv("MAIL_SERVER")
app.config["MAIL_PORT"] = int(os.getenv("MAIL_PORT"))  # Convert to int
app.config["MAIL_USE_TLS"] = os.getenv("MAIL_USE_TLS").lower() in ["true", "1"]
app.config["MAIL_USE_SSL"] = os.getenv("MAIL_USE_SSL").lower() in ["true", "1"]
app.config["MAIL_USERNAME"] = os.getenv("MAIL_USERNAME")
app.config["MAIL_PASSWORD"] = os.getenv("MAIL_PASSWORD")
app.config["MAIL_DEFAULT_SENDER"] = os.getenv("MAIL_DEFAULT_SENDER")

app.debug = True

# Initialize database & Flask extensions
db.init_app(app)
#migrate = Migrate(app, db)  # Initialize Flask-Migrate
mail = Mail(app)  # Initialize Flask-Mail
login_manager = LoginManager(app)
login_manager.login_view = 'login'  # Redirect to login if user isn't authenticated

@login_manager.user_loader
def load_user(user_id):
    """ Load user session for Flask-Login """
    try:
        return db.session.get(User, int(user_id)) if user_id else None
    except ValueError:
        return None

# Create database tables if they don't exist
with app.app_context():
    db.create_all()

# Routes

# Nav-Bar
@app.route('/editor')
def editor():
    return render_template('editor.html')

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/dashboard')
@login_required
def dashboard():
    resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.uploaded_at.desc()).all()
    return render_template('dashboard.html', resumes=resumes)

@app.route('/analyze', methods=['GET', 'POST'])
@login_required
def analyze_resume_route():
    if request.method == 'POST':
        resume_file = request.files.get('resume')
        job_desc = request.form.get('job_description', '')
        region = request.form.get('region', '')
        country = request.form.get('country', '')

        # Validate inputs
        if not resume_file or resume_file.filename == '':
            flash('No file selected. Please upload a resume.', 'error')
            return redirect(request.url)

        if not region or not country:
            flash('Please select both a region and a country.', 'error')
            return redirect(request.url)

        try:
            # Determine file extension
            file_extension = resume_file.filename.split('.')[-1].lower()
            resume_text = ""

            if file_extension == 'docx':
                # Read DOCX file and extract text
                doc = docx.Document(resume_file)
                resume_text = "\n".join([para.text for para in doc.paragraphs])

            elif file_extension == 'pdf':
                # Read PDF file and extract text
                pdf_reader = PyPDF2.PdfReader(resume_file)
                resume_text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])

                if not resume_text.strip():  # Handle scanned PDFs
                    flash("This PDF file may be scanned or non-selectable. Please upload a DOCX instead.", "error")
                    return redirect(request.url)

            else:
                flash("Unsupported file type. Please upload a PDF or DOCX file.", "error")
                return redirect(request.url)

            # Upload file to AWS S3
            s3_path = upload_to_s3(resume_file, current_user.id, region, country)

            if not s3_path:
                flash("Error uploading file to S3. Please try again.", "error")
                return redirect(request.url)

            # AI Resume Analysis
            analysis = analyze_resume(resume_text, job_desc)

            # Save analysis & file metadata to database
            new_resume = Resume(
                user_id=current_user.id,
                s3_path=s3_path,
                region=region,
                country=country,
                analysis=analysis,
                uploaded_at=datetime.now(timezone.utc)  # Ensure timezone-aware datetime
            )
            db.session.add(new_resume)
            db.session.commit()
            fine_tuner(resume_text, job_desc)
            flash("Resume successfully uploaded and analyzed!", "success")
            return render_template('resume/analysis.html', analysis=analysis)

        except Exception as e:
            db.session.rollback()  # Rollback transaction in case of an error
            flash(f'Error processing file: {str(e)}', 'error')
            return redirect(url_for('analyze_resume_route'))

    return render_template('resume/analyze.html')

@app.route('/history')
@login_required
def user_history():
    # Get user's uploaded resumes
    resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.uploaded_at.desc()).all()
    return render_template('resume/history.html', resumes=resumes)

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email')
        user = User.query.filter_by(email=email).first()

        if user:
            token = generate_timed_token(user.email)
            reset_url = url_for('reset_password', token=token, _external=True)
            html = render_template("auth/reset_password_email.html", reset_url=reset_url)
            send_email(user.email, "Password Reset Request", html)

        flash("If your email exists, you will receive a password reset link.", "success")
        return redirect(url_for('login'))
    
    return render_template('auth/forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    email = confirm_timed_token(token)

    if not email:
        flash('Invalid or expired token', 'error')
        return redirect(url_for('login'))

    if request.method == 'POST':
        new_password = request.form.get('password')
        user = User.query.filter_by(email=email).first()

        if user:
            user.set_password(new_password)  # Hash new password
            db.session.commit()
            flash('Password reset successful. Please log in.', 'success')
            return redirect(url_for('login'))

    return render_template('auth/reset_password.html', token=token)

@app.route('/interview', methods=['GET', 'POST'])
@login_required
def interview():
    if request.method == 'POST':
        answer = request.form.get('answer', '')
        if not answer:
            flash('Please provide an answer', 'error')
            return redirect(url_for('interview'))
            
        try:
            evaluation = evaluate_star_response(answer)
            return render_template('evaluation.html', evaluation=evaluation)
        except Exception as e:
            flash(f'Evaluation error: {str(e)}', 'error')
            return redirect(url_for('interview'))
    
    try:
        question = generate_interview_question("Technology", "Mid-Level")
    except Exception as e:
        question = f"Error generating question: {str(e)}"
        flash(question, 'error')
    
    return render_template('interview/interview.html', question=question)

# Authentication routes
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        user = User.query.filter_by(email=email).first()

        if user and user.check_password(password):  # Verify password hash
            login_user(user)
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))

        flash('Invalid email or password', 'error')

    return render_template('auth/login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        if User.query.filter_by(email=email).first():
            flash('Email already exists', 'error')
            return redirect(url_for('register'))

        # Create new user and hash password
        new_user = User(email=email)
        new_user.set_password(password)  # Hash the password

        # Save to database
        db.session.add(new_user)
        db.session.commit()

        flash('Registration successful! Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template('auth/register.html')


@app.route('/resume/<int:resume_id>')
@login_required
def view_resume(resume_id):
    """Generate a temporary S3 pre-signed URL for viewing the resume"""
    resume = Resume.query.get_or_404(resume_id)

    # Generate pre-signed URL for the stored S3 file key
    presigned_url = generate_presigned_url(resume.s3_path)

    if presigned_url:
        logging.info(f"Redirecting to pre-signed URL: {presigned_url}")
        return redirect(presigned_url)
    else:
        logging.error("Error generating pre-signed URL.")
        flash("Error generating pre-signed URL.", "error")
        return redirect(url_for("user_history"))
    
@app.route('/generate_resume', methods=['POST'])
def generate_resume_route():
    try:
        # Ensure the data is being received correctly
        user_answer = request.json.get('answer')
        existing_resume = request.json.get('existing_resume', '')  # Get existing resume content
        
        if not user_answer:
            return jsonify({'error': 'No input provided'}), 400
        
        # Merge existing resume with new input
        combined_input = existing_resume + "\n" + user_answer if existing_resume else user_answer
        
        # Call the AI function to update the resume
        updated_resume = generate_resume(combined_input)
        
        return jsonify({'resume': updated_resume})
    
    except Exception as e:
        print("Error generating resume:", str(e))  # Debugging log
        return jsonify({'error': str(e)}), 400

@app.route('/export_resume', methods=['POST'])
def export_resume():
    try:
        # Get the resume text from the request
        resume_text = request.json.get('resume')
        
        if not resume_text:
            return jsonify({'error': 'No resume text provided'}), 400
        
         # Create a Word document
        doc = Document()
        doc.add_heading('Generated Resume', level=1)
        doc.add_paragraph(resume_text)

        # Save the document in memory
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Return the file for download
        return send_file(file_stream, as_attachment=True, download_name="resume.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document

    except Exception as e:
        return jsonify({'error': str(e)}), 400


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('home'))

# Error handling
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template('500.html'), 500

# Run Flask app
if __name__ == '__main__':
    app.run(debug=False)
